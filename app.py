from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
import os
import tempfile
from docx import Document
import openai

app = Flask(__name__)
CORS(app)

# OpenAI API кілті (жергілікті немесе .env арқылы)
openai.api_key = os.environ.get("OPENAI_API_KEY")

# GPT үшін prompt құру
def build_glossary_prompt(topic, language_level):
    return f"""
You are a CLIL glossary generator for Kazakhstani Informatics teachers.

Generate a glossary of 15 English terms related to the topic "{topic}" suitable for students with {language_level} English proficiency.

For each term, provide:
1. Term (in English)
2. Translation in Kazakh
3. Translation in Russian
4. IPA pronunciation (e.g., /ˈælgərɪðəm/)
5. A simple phonetic Latin-style transcription of how to pronounce the word (e.g., 'al-guh-rith-um')
6. A very short English definition (max 15 words)
7. A very short Kazakh translation of that definition (max 15 words)

Return the glossary as a markdown table like this:
| Term | Kazakh | Russian | IPA | How to Read | Definition |
Each row’s definition column must contain:

**EN:** English definition  
**KZ:** Қазақша аудармасы
"""

@app.route("/generate_glossary", methods=["POST"])
def generate_glossary():
    data = request.json
    topic = data.get("topic", "")
    level = data.get("language_level", "B1-B2")

    prompt = build_glossary_prompt(topic, level)

    try:
        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",
            messages=[
                {"role": "system", "content": "You are a professional CLIL glossary generator."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.7
        )

        result = response.choices[0].message["content"]
        return jsonify({"glossary": result})

    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/download_glossary_docx", methods=["POST"])
def download_docx():
    data = request.json
    markdown = data.get("glossary", "")

    if not markdown:
        return jsonify({"error": "Glossary is empty"}), 400

    # DOCX жасау
    doc = Document()
    doc.add_heading("CLIL Glossary", level=1)

    lines = markdown.strip().split('\n')
    for line in lines:
        doc.add_paragraph(line)

    # Уақытша файл жасау
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(temp_file.name)

    return send_file(temp_file.name, as_attachment=True, download_name="glossary.docx")

if __name__ == "__main__":
    app.run(debug=True)
